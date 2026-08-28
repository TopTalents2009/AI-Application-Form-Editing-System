"""把修改对照表写成 Word（.docx），与 Markdown 对照表字段对齐。"""
from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HEADERS = ["#", "章节", "意见条款", "改前摘录", "修改意见", "改后摘录", "结果"]
COL_CM = (1.2, 2.0, 4.4, 5.2, 5.2, 5.2, 2.6)


def _east_asia(run, name="微软雅黑", size=9, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)


def _shade(cell, fill: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _cell_text(cell, text, *, bold=False, size=9, color=None, fill=None, center=False):
    if fill:
        _shade(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text or "").strip())
    _east_asia(run, size=size, bold=bold, color=color)


def _set_repeat_header(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def _status_style(st: str):
    s = str(st or "")
    if "未命中" in s:
        return (197, 34, 31), "FDE9E8"
    if "未检出" in s:
        return (178, 106, 0), "FDF3DF"
    if "已改" in s:
        return (24, 128, 56), "E2F5E7"
    return (91, 101, 119), "EEF1F6"


def write_compare_docx(path, *, app_name: str, app_no: str, created: str, rows: list, leftovers: list):
    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = Cm(29.7), Cm(21.0)
    sec.left_margin = Cm(1.2)
    sec.right_margin = Cm(1.2)
    sec.top_margin = Cm(1.3)
    sec.bottom_margin = Cm(1.3)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("修改对照表")
    _east_asia(r, size=16, bold=True)

    sub = doc.add_paragraph()
    meta = "申报书编号 " + (app_no or "未识别") + "　" + (app_name or "") + "　生成时间 " + (created or "")
    r2 = sub.add_run(meta)
    _east_asia(r2, size=10, color=(91, 101, 119))
    note = doc.add_paragraph()
    r3 = note.add_run("管线：大模型出计划 → 人工修订 → 内置执行器落盘。对照表供核对，不替代申报书正文。")
    _east_asia(r3, size=9, color=(91, 101, 119))

    table = doc.add_table(rows=1, cols=len(HEADERS))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"

    hdr = table.rows[0]
    _set_repeat_header(hdr)
    for i, h in enumerate(HEADERS):
        hdr.cells[i].width = Cm(COL_CM[i])
        _cell_text(hdr.cells[i], h, bold=True, size=9, fill="1F3B73", color=(255, 255, 255), center=True)

    data = list(rows or [])
    if not data:
        data = [{"n": "", "section": "", "clause": "（无编辑行）", "find": "", "opinion": "", "replace": "", "status": ""}]
    for i, item in enumerate(data):
        tr = table.add_row()
        st = str(item.get("status") or "")
        sc, sf = _status_style(st)
        vals = [
            str(item.get("n") or (i + 1 if st else "")),
            item.get("section") or "",
            item.get("clause") or "",
            item.get("find") or "",
            item.get("opinion") or "",
            item.get("replace") or "",
            st,
        ]
        for j, val in enumerate(vals):
            tr.cells[j].width = Cm(COL_CM[j])
            center = j in (0, 1, 6)
            if j == 6 and st:
                _cell_text(tr.cells[j], val, size=9, bold=True, color=sc, fill=sf, center=True)
            else:
                _cell_text(tr.cells[j], val, size=9, center=center)

    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(14)
    r4 = h2.add_run("遗留事项（需人工补充真实数据）")
    _east_asia(r4, size=13, bold=True)
    if leftovers:
        for i, line in enumerate(leftovers, 1):
            p = doc.add_paragraph()
            rr = p.add_run(str(i) + ". " + str(line))
            _east_asia(rr, size=10)
    else:
        p = doc.add_paragraph()
        rr = p.add_run("（无）")
        _east_asia(rr, size=10, color=(152, 161, 179))

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return Path(path)
