# -*- coding: utf-8 -*-
"""申报书成品完整性校验：Word 能打开并统计内容；Excel 能打开并统计非空行。
用法: python sb_verify.py <文件>
输出: OK ... ；失败非零退出。
"""
import os
import shutil
import sys
import tempfile

WORD_EXT = {".docx", ".docm", ".wps"}
EXCEL_EXT = {".xlsx", ".xlsm", ".xls"}


def _ext(path):
    return os.path.splitext(path)[1].lower()


def _verify_word(path):
    import docx
    from apply_edits import _set_word_main_ct

    ext = _ext(path)
    work = path
    tmp = None
    try:
        if ext == ".docm":
            fd, tmp = tempfile.mkstemp(suffix=".docx")
            os.close(fd)
            shutil.copyfile(path, tmp)
            _set_word_main_ct(tmp, macro=False)
            work = tmp
        d = docx.Document(work)
        paras = sum(1 for p in d.paragraphs if p.text.strip())
        tables = len(d.tables)
    finally:
        if tmp and os.path.exists(tmp):
            try:
                os.remove(tmp)
            except OSError:
                pass
    print("OK paras=%d tables=%d" % (paras, tables))
    return 0


def _verify_excel(path):
    ext = _ext(path)
    if ext in (".xlsx", ".xlsm"):
        from openpyxl import load_workbook

        wb = load_workbook(path, read_only=True, data_only=True, keep_vba=False)
        try:
            n_sheets = len(wb.worksheets)
            n_rows = 0
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    if any(c is not None and str(c).strip() for c in row):
                        n_rows += 1
        finally:
            wb.close()
        print("OK sheets=%d rows=%d" % (n_sheets, n_rows))
        return 0
    import xlrd

    book = xlrd.open_workbook(str(path))
    n_rows = 0
    for sheet in book.sheets():
        for r in range(sheet.nrows):
            if any(str(sheet.cell_value(r, c)).strip() for c in range(sheet.ncols)):
                n_rows += 1
    print("OK sheets=%d rows=%d" % (book.nsheets, n_rows))
    return 0


def main(argv):
    if len(argv) < 2:
        print("usage: sb_verify.py <file>", file=sys.stderr)
        return 2
    path = argv[1]
    ext = _ext(path)
    here = os.path.dirname(os.path.abspath(__file__))
    if here not in sys.path:
        sys.path.insert(0, here)
    try:
        if ext in WORD_EXT:
            return _verify_word(path)
        if ext in EXCEL_EXT:
            return _verify_excel(path)
        print("BROKEN unsupported %s" % ext, file=sys.stderr)
        return 1
    except Exception as e:
        print("BROKEN %s" % e, file=sys.stderr)
        return 1


if __name__ == "__main__":
    sys.exit(main(sys.argv))
